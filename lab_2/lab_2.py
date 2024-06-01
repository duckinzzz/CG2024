from docx import Document
from docx.shared import RGBColor

doc = Document()

heading = doc.add_heading('Повестка на свадьбу')
heading_font = heading.style.font
heading_font.color.rgb = RGBColor(255, 0, 0)  # Например, красный цвет

table = doc.add_table(rows=2, cols=2)
table.rows[0].cells[0].text = 'Гражданину '
table.rows[0].cells[1].text = '{{ initials }}'
table.rows[1].cells[0].text = 'Проживающему '
table.rows[1].cells[1].text = '{{ conscript_address }}'

p = doc.add_paragraph('')
p.add_run(
    'На основании Дружеского Закона ')
p.add_run('"О свадебной обязанности" ').italic = True
p.add_run(
    'Вы подлежите призыву на свадебное торжество и обязаны {{ date }} к {{ time }} часам '
    'явиться по адресу: {{ wedding_address }} для ')
p.add_run('проведения мероприятий связанных с венчанием молодоженов.').bold = True

doc.add_heading('При себе иметь: ', 1)
doc.add_paragraph('Хорошее настроение',
                  style='List Bullet')
doc.add_paragraph('Костюм',
                  style='List Bullet')
doc.add_paragraph('Подарок',
                  style='List Bullet')

p = doc.add_paragraph('Подпись {{ signature }}')

p = doc.add_paragraph('Печать {{ stamp }}')

doc.save('lab2_tmp.docx')
